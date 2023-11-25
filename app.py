# Programmers: Anita Martin, Chris Heise, Dion Boyer, and Felix Jaramillo
# Course: BSSD 4350 - Agile Methodologies
# Instructor: Jonathan Lee
# Program: Inclusivity Editor App
# File: app.py

import gradio as gr
from difflib import Differ
from docx import *
from fpdf import *
import pyperclip
import together

# TODO: get rid of functions we don't need and unused variables

together.api_key = "YOUR API KEY HERE"

users_text = ""

# For testing, will be what the user inputs
<<<<<<< Updated upstream
EXAMPLE_TEXT = """Lorem ipsum dolor sit amet, consectetur adipiscing elit. Mauris ultricies elementum nulla, id placerat nunc efficitur non. Nam tempus, nulla ac sodales laoreet, lacus tellus efficitur enim, eget sodales lorem ante ut neque. Mauris quis eros sed velit mollis porta. Aenean libero diam, sagittis sed arcu non, fermentum tincidunt leo. Nulla sed velit tempor, dapibus ex in, rhoncus orci. Praesent sit amet odio sagittis arcu venenatis consequat vitae vitae tortor. Sed et maximus nunc, nec placerat ligula.
Etiam libero nisi, fringilla a imperdiet in, luctus a tortor. Pellentesque quis venenatis velit, quis malesuada nisl. Praesent eu placerat ante. Vivamus quis mi porttitor, faucibus purus non, tempus lacus. Pellentesque et imperdiet dui. Vivamus ut lacus quis lacus maximus iaculis. Vivamus mollis odio orci, ut egestas nisl rhoncus nec. Quisque sit amet lorem viverra, lobortis erat quis, consectetur augue. Etiam blandit tempus purus nec maximus. Vestibulum tempus semper ipsum ac faucibus."""

# For testing, will be what the LLM returns
CORRECTED_TEXT = """Lorem ipsum dolor sit amet, consectetur adipiscing elit. Mauris ultricies elementum nulla, id placerat nunc efficitur non. Nulla ac sodales laoreet, lacus tellus efficitur enim, eget sodales lorem ante ut neque. Mauris quis eros sed velit mollis porta. Aenean libero diam, sagittis sed arcu non, fermentum tincidunt leo. Nulla sed velit tempor, dapibus ex in, rhoncus orci. Praesent sit amet odio sagittis arcu venenatis consequat vitae vitae tortor. Sed et maximus nunc, nec placerat ligula.
Etiam libero nisi, fringilla a imperdiet in. Pellentesque quis venenatis velit, elit malesuada nisl. Praesent eu placerat ante. Vivamus quis mi porttitor, faucibus purus non, tempus lacus. Pellentesque et imperdiet dui. Vivamus ut lacus quis lacus maximus iaculis. Vivamus mollis odio orci, ut egestas nisl rhoncus nec. Quisque sit amet lorem viverra, lobortis erat quis, consectetur augue. Etiam blandit tempus purus nec maximus. Vestibulum tempus semper ipsum ac tortor."""
=======
EXAMPLE_TEXT = ""

# For testing, will be what the LLM returns
CORRECTED_TEXT = ""
>>>>>>> Stashed changes

textInput = ""

DIFFERENCES = []

# Global Components (accessed by multiple tabs/pages)
original_text = gr.Textbox(
    label="Your Text",
    info="Your original text.",
    lines=10,
)

previewText = gr.Textbox(label="Output Textbox", value=textInput)

# Isn't currently working. Seems to need to be called with a button click like other componenets/functions
# Source: https://github.com/gradio-app/gradio/issues/2412


def update_preview(text):
    gr.Textbox(label="Output Textbox", value=textInput)
    return text


def change_page(page_number):
    """Changes the page to the page number passed in."""
    return gr.Tabs.update(selected=page_number)


def download(output, type):
    match type:
        case "PDF":
            """Download text as a PDF."""
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            pdf.multi_cell(0, 10, str(output))
            pdf_file = "history_download.pdf"
            pdf.output(pdf_file)
            return pdf_file
        case "DOCX":
            doc = Document()

            doc.add_heading('History Download', 0)
            doc.add_paragraph(output)
            doc_output = "history_download.docx"
            doc.save(doc_output)

            return doc_output
        case "TXT":
            text_file = 'history_download.txt'
            with open('history_download.txt', 'w') as txt_file:
                txt_file.write(output)
                txt_file.close()
            return text_file
        case _:
            """Download text as a PDF."""
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            pdf.multi_cell(0, 10, str(output))
            pdf_file = type + "history_download.pdf"
            pdf.output(pdf_file)
            return pdf_file


def copy_text(output):
    """Copy text to the clipboard."""
    pyperclip.copy(output)
    text = pyperclip.paste()


def load_text(temp_file):
    """Load text from a temporary file."""
    content = ""
    with open(temp_file.name, "r", encoding="utf-8") as f:
        content = f.read()
    return content


def submit_text(text):
    users_text = text
    text_input += users_text
    change_page(2)
    return users_text


def diff_texts(text1, text2):
    """Find the differences between two texts."""
    d = Differ()
    return [
        (token[2:], token[0] if token[0] != " " else None)
        for token in d.compare(text1, text2)
    ]
<<<<<<< Updated upstream
=======


def call_llm(prompt_text):
    llm = together.Complete.create(
        # TODO: Add the selected preferences choice here to the string instead of hard coding it in.
        prompt=' Rewrite as a Professional Correspondence'+ prompt_text,
        model="togethercomputer/llama-2-7b-chat",
        max_tokens=256,
        temperature=0.8,
        top_k=60,
        top_p=0.6,
        repetition_penalty=1.1,
        stop=['<human>']
    )
    # print(llm['prompt'])
    # print(llm['output']['choices'][0]['text'])
    answer = (llm['output']['choices'][0]['text']).strip().split("Answer:\n")[0]
    return answer


>>>>>>> Stashed changes
with gr.Blocks() as incluesive:
    gr.Markdown("# INCLUeSIVE")
    with gr.Tabs() as pages:
        """FIRST PAGE"""
        with gr.TabItem("Welcome", id=0) as first_page:
            with gr.Tabs():
                with gr.TabItem("Directions"):
                    gr.Markdown(
                        "Welcome to Incluesive an app that will help correcct your writings to be more incluesive of everyone. "
                        "To use Incluesive Pick a writing purpose then enter your text into the text box and submit. After you submit the changes to your text will be shown.")
                with gr.TabItem("Preferences"):
                    pref = gr.Button(value="Preferences", size='sm')
                    preference_choice = gr.Radio(
                        ["Professional Correspondence", "Personal Correspondence", "Educational Paper",
                         "Technical Instructions"], label="Writing purpose")
                    submit_button = gr.Button("Submit", link="")
                    submit_button.click(inputs=preference_choice, outputs=None)
        """END FIRST PAGE"""

        """SECOND PAGE"""
        with gr.TabItem("Input", id=1) as second_page:
            with gr.Tab("Type/Paste"):
                text_input = gr.Textbox(
                    label="Type or paste your text here.",
                    info="Your Original Text.",
                    lines=10,
                    value=EXAMPLE_TEXT,
                )
                submit_button = gr.Button("Submit")
                corrected_text = gr.Textbox(
                    label="Corrected Text",
                    info="Our suggested corrected text",
                    lines=10,
                    value=CORRECTED_TEXT,
                    visible=False,
                )
                submit_button.click(submit_text, inputs=[text_input], outputs=original_text)
<<<<<<< Updated upstream
                clear_button = gr.ClearButton()
                text_input.change(submit_text, text_input, original_text)
=======
                clear_button = gr.ClearButton(text_input)
>>>>>>> Stashed changes
            with gr.Tab("Upload"):
                file_input = gr.File(
                    file_types=["text"],
                )
                with gr.Row():
                    upload_button = gr.Button("Upload")
                loaded_text = gr.Textbox(
                    label="Your Text",
                    info="The text you uploaded.",
                    lines=10,
                )
                corrected_text = gr.Textbox(
                    label="Corrected Text",
                    info="Our suggested corrected text",
                    lines=10,
                    value=CORRECTED_TEXT,
                    visible=False,
                )
                with gr.Row():
                    clear_button = gr.ClearButton(loaded_text)
                    submit_button = gr.Button("Submit")
                    upload_button.click(load_text, inputs=[file_input], outputs=[loaded_text])
                    submit_button.click(submit_text, inputs=[loaded_text], outputs=original_text)
        """END SECOND PAGE"""

        """THIRD PAGE"""
        with gr.TabItem("Results", id=2) as third_page:
<<<<<<< Updated upstream
            original_text.render()
            original_text.change(update_preview, original_text, previewText)
=======
            input_text = original_text.render()
            output_text = gr.Textbox(label="Results from LLM")
>>>>>>> Stashed changes
            with gr.Row():
                submit_button = gr.Button("Submit")
                clear_button = gr.ClearButton(original_text)

            submit_button.click(fn=call_llm, inputs=input_text, outputs=output_text)

            with gr.Row():
                # TODO: Get buttons to do what they are suppose to do inside the TextBox Results from LLM
                submit_paragraph_button = gr.Button("Accept")
                accept_paragraph_button = gr.Button("Ignore")
                done_paragraph_button = gr.Button("Done")
        """END THIRD PAGE"""

        """FOURTH PAGE"""
<<<<<<< Updated upstream
        with gr.TabItem("Save", id=3) as fourth_page:
=======
        # TODO: Pass the data from the 3rd page TextBox Results from LLM so we can save
        with gr.TabItem("Save", id=3) as third_page:
>>>>>>> Stashed changes
            with gr.Accordion(label="Account"):
                preferences = gr.Button(value="Preferences")
                signout = gr.Button(value="Sign Out")
            with gr.Row():
                previewText.render()
                file = gr.File()
            with gr.Row():
                dropdown_type = gr.Dropdown(
                    ["DOCX", "PDF", "TXT"], label="File Type", info="Select your file type."
                )
                download_btn = gr.Button(value="Download", scale=0)
                copy_btn = gr.Button(value="Copy", scale=0)
                done_btn = gr.Button(value="Done", scale=0)
                download_btn.click(fn=download, inputs=[previewText, dropdown_type], outputs=file, api_name="Download")
                copy_btn.click(fn=copy_text, inputs=previewText, api_name="Copy")
        """END FOURTH PAGE"""
<<<<<<< Updated upstream
incluesive.launch()
=======
#call_llm()
incluesive.launch()
>>>>>>> Stashed changes
